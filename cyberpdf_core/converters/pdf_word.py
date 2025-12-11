"""
PDF <-> Word conversion
"""
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from pathlib import Path
import subprocess
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class PDFWordConverter:
    """Convert between PDF and Word formats"""
    
    @staticmethod
    def pdf_to_word(input_path: str, output_path: str, method: str = "auto") -> str:
        """
        Convert PDF to Word document
        
        Args:
            input_path: Path to input PDF
            output_path: Path for output DOCX
            method: Conversion method ('auto', 'text', 'libreoffice')
        
        Returns:
            Path to output DOCX
        """
        if method == "auto":
            # Try LibreOffice first, fall back to text extraction
            try:
                return PDFWordConverter._pdf_to_word_libreoffice(input_path, output_path)
            except Exception as e:
                logger.warning(f"LibreOffice conversion failed: {e}, falling back to text extraction")
                return PDFWordConverter._pdf_to_word_text(input_path, output_path)
        
        elif method == "libreoffice":
            return PDFWordConverter._pdf_to_word_libreoffice(input_path, output_path)
        
        elif method == "text":
            return PDFWordConverter._pdf_to_word_text(input_path, output_path)
        
        else:
            raise ValueError(f"Unknown conversion method: {method}")
    
    @staticmethod
    def _pdf_to_word_libreoffice(input_path: str, output_path: str) -> str:
        """Convert PDF to Word using LibreOffice (best quality)"""
        output_dir = Path(output_path).parent
        
        # Use LibreOffice in headless mode
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to", "docx",
                "--outdir", str(output_dir),
                input_path
            ],
            capture_output=True,
            text=True,
            timeout=60
        )
        
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
        
        # LibreOffice creates file with same name but .docx extension
        generated_file = output_dir / f"{Path(input_path).stem}.docx"
        
        # Rename if needed
        if str(generated_file) != output_path:
            generated_file.rename(output_path)
        
        logger.info(f"Converted PDF to Word using LibreOffice: {output_path}")
        return output_path
    
    @staticmethod
    def _pdf_to_word_text(input_path: str, output_path: str) -> str:
        """Convert PDF to Word by extracting text (fallback method)"""
        doc = fitz.open(input_path)
        word_doc = Document()
        
        # Add title from metadata
        metadata = doc.metadata
        if metadata.get("title"):
            word_doc.add_heading(metadata["title"], 0)
        
        # Extract text from each page
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            
            if text.strip():
                # Add page content
                word_doc.add_paragraph(text)
                
                # Add page break except for last page
                if page_num < len(doc) - 1:
                    word_doc.add_page_break()
        
        word_doc.save(output_path)
        doc.close()
        
        logger.info(f"Converted PDF to Word using text extraction: {output_path}")
        return output_path
    
    @staticmethod
    def word_to_pdf(input_path: str, output_path: str, method: str = "auto") -> str:
        """
        Convert Word document to PDF
        
        Args:
            input_path: Path to input DOCX
            output_path: Path for output PDF
            method: Conversion method ('auto', 'libreoffice', 'unoconv')
        
        Returns:
            Path to output PDF
        """
        if method == "auto":
            # Try LibreOffice first
            try:
                return PDFWordConverter._word_to_pdf_libreoffice(input_path, output_path)
            except Exception as e:
                logger.warning(f"LibreOffice conversion failed: {e}")
                raise RuntimeError("Word to PDF conversion requires LibreOffice")
        
        elif method == "libreoffice":
            return PDFWordConverter._word_to_pdf_libreoffice(input_path, output_path)
        
        elif method == "unoconv":
            return PDFWordConverter._word_to_pdf_unoconv(input_path, output_path)
        
        else:
            raise ValueError(f"Unknown conversion method: {method}")
    
    @staticmethod
    def _word_to_pdf_libreoffice(input_path: str, output_path: str) -> str:
        """Convert Word to PDF using LibreOffice"""
        output_dir = Path(output_path).parent
        
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(output_dir),
                input_path
            ],
            capture_output=True,
            text=True,
            timeout=60
        )
        
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
        
        # LibreOffice creates file with same name but .pdf extension
        generated_file = output_dir / f"{Path(input_path).stem}.pdf"
        
        # Rename if needed
        if str(generated_file) != output_path:
            generated_file.rename(output_path)
        
        logger.info(f"Converted Word to PDF using LibreOffice: {output_path}")
        return output_path
    
    @staticmethod
    def _word_to_pdf_unoconv(input_path: str, output_path: str) -> str:
        """Convert Word to PDF using unoconv"""
        result = subprocess.run(
            ["unoconv", "-f", "pdf", "-o", output_path, input_path],
            capture_output=True,
            text=True,
            timeout=60
        )
        
        if result.returncode != 0:
            raise RuntimeError(f"unoconv conversion failed: {result.stderr}")
        
        logger.info(f"Converted Word to PDF using unoconv: {output_path}")
        return output_path
    
    @staticmethod
    def check_dependencies() -> dict:
        """
        Check which conversion tools are available
        
        Returns:
            Dictionary with availability status
        """
        dependencies = {
            "libreoffice": False,
            "unoconv": False,
        }
        
        # Check LibreOffice
        try:
            result = subprocess.run(
                ["libreoffice", "--version"],
                capture_output=True,
                timeout=5
            )
            dependencies["libreoffice"] = result.returncode == 0
        except:
            pass
        
        # Check unoconv
        try:
            result = subprocess.run(
                ["unoconv", "--version"],
                capture_output=True,
                timeout=5
            )
            dependencies["unoconv"] = result.returncode == 0
        except:
            pass
        
        return dependencies
